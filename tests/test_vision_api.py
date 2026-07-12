from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient
from openpyxl import Workbook

from backend.app.core.settings import AppSettings
from backend.app.main import create_app


def _client(tmp_path, monkeypatch) -> TestClient:
    settings = AppSettings(vision_upload_directory=tmp_path / "documents")
    monkeypatch.setattr("backend.app.api.routes.vision.get_settings", lambda: settings)
    return TestClient(create_app(settings))


def test_document_upload_summary_question_and_table_extraction(
    tmp_path, monkeypatch
) -> None:
    document = Document()
    document.add_heading("July invoice", 0)
    document.add_paragraph(
        "The invoice amount is 1,250 dollars and payment is due on Friday."
    )
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "Amount"
    table.cell(1, 0).text = "Design"
    table.cell(1, 1).text = "$1,250"
    payload = BytesIO()
    document.save(payload)

    client = _client(tmp_path, monkeypatch)
    response = client.post(
        "/api/v1/documents",
        files={
            "file": (
                "invoice.docx",
                payload.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    record = response.json()["data"]
    assert record["document_type"] == "docx"
    assert record["tables"][0][1][1] == "$1,250"

    tables = client.get(f"/api/v1/documents/{record['id']}/tables")
    assert tables.status_code == 200
    assert tables.json()["data"][0][0] == ["Item", "Amount"]

    summary = client.post(f"/api/v1/documents/{record['id']}/summarize")
    assert summary.status_code == 200
    assert "invoice amount" in summary.json()["data"]["summary"]

    answer = client.post(
        f"/api/v1/documents/{record['id']}/questions",
        json={"question": "What is the invoice amount?"},
    )
    assert answer.status_code == 200
    assert "1,250" in answer.json()["data"]["answer"]


def test_document_upload_rejects_an_unsupported_format(tmp_path, monkeypatch) -> None:
    client = _client(tmp_path, monkeypatch)

    response = client.post(
        "/api/v1/documents",
        files={"file": ("unsafe.exe", b"data", "application/octet-stream")},
    )

    assert response.status_code == 415


def test_excel_upload_preserves_sheet_rows(tmp_path, monkeypatch) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Budget"
    sheet.append(["Category", "Amount"])
    sheet.append(["Hosting", 45])
    payload = BytesIO()
    workbook.save(payload)
    client = _client(tmp_path, monkeypatch)

    response = client.post(
        "/api/v1/documents",
        files={
            "file": (
                "budget.xlsx",
                payload.getvalue(),
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        },
    )

    assert response.status_code == 200
    data = response.json()["data"]
    assert data["sheet_names"] == ["Budget"]
    assert data["tables"][0][1] == ["Hosting", "45"]


def test_vision_analysis_reports_text_buttons_and_errors(tmp_path, monkeypatch) -> None:
    settings = AppSettings(vision_upload_directory=tmp_path / "documents")
    monkeypatch.setattr(
        "backend.app.vision.vision_service.VisionService.detect_tesseract_command",
        staticmethod(lambda settings: "tesseract"),
    )
    monkeypatch.setattr(
        "backend.app.vision.vision_service.pytesseract.image_to_data",
        lambda image, output_type: {
            "text": ["Error", "Save", "failed"],
            "conf": ["95", "91", "88"],
            "left": [1, 50, 100],
            "top": [2, 3, 4],
            "width": [30, 25, 30],
            "height": [10, 10, 10],
        },
    )
    image_bytes = BytesIO()
    from PIL import Image

    Image.new("RGB", (160, 90), "white").save(image_bytes, format="PNG")
    monkeypatch.setattr("backend.app.api.routes.vision.get_settings", lambda: settings)
    client = TestClient(create_app(settings))

    response = client.post(
        "/api/v1/vision/analyze",
        files={"file": ("error.png", image_bytes.getvalue(), "image/png")},
    )

    assert response.status_code == 200
    data = response.json()["data"]
    assert data["text"] == "Error Save failed"
    assert data["ui_elements"][0]["label"] == "Save"
    assert data["errors"]
